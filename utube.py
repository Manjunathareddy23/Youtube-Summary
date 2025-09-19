# ---------------- Part 1: Imports and Setup ----------------
import streamlit as st
import os
import googleapiclient.discovery
from dotenv import load_dotenv
import google.generativeai as genai
from youtube_transcript_api import YouTubeTranscriptApi, NoTranscriptFound, TranscriptsDisabled
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
import time
from urllib.parse import urlparse, parse_qs
from concurrent.futures import ThreadPoolExecutor, as_completed

# Constants
RETRY_COUNT = 3
RETRY_DELAY = 2

# Streamlit page config
st.set_page_config(
    page_title="Gemini Flash YouTube Video Summary App",
    page_icon="🎯",
    layout="wide"
)

# ---------------- Session State ----------------
session_defaults = {
    "current_summary": None,
    "word_doc_binary": None,
    "video_processed": False,
    "current_transcript": None,
    "current_video_id": None,
    "current_video_title": None,
    "qa_history": [],
    "clear_input": False,
    "video_count": 0,
    "query_count": 0,
    "fast_summary_generated": False,
}
for key, value in session_defaults.items():
    if key not in st.session_state:
        st.session_state[key] = value

# ---------------- Load API Keys ----------------
load_dotenv(override=True)
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# ---------------- Prompts ----------------
CHUNK_PROMPT = """Analyze this portion of the video transcript and provide:
1. Key points discussed in this section
2. Notable quotes with timestamps
3. Technical data or statistics
4. Important concepts and definitions
5. Brief summary

Transcript section: """

FINAL_PROMPT = """Based on the analysis of all sections, provide a comprehensive summary with:
1. Main Topic/Title
2. Executive Summary (200 words)
3. Key Points (10-20)
4. Detailed Analysis (2000 words)
5. Notable Quotes and Key Statements
6. Technical Data & Statistics
7. Key Terms & Definitions
8. Concepts & Frameworks
9. Timeline & Structure
10. Practical Applications

Please synthesize this complete summary: """

FAST_SUMMARY_PROMPT = """Provide a concise executive summary of the video transcript (200-500 words). Skip detailed quotes or technical terms."""

# ---------------- Helper Functions ----------------
def get_youtube_video_id(url):
    """Extract video ID from various YouTube URL formats."""
    url = url.strip()
    if not url:
        return None
    patterns = [
        r'(?:https?:\/\/)?(?:www\.)?youtube\.com\/watch\?v=([a-zA-Z0-9_-]+)',
        r'(?:https?:\/\/)?(?:www\.)?m\.youtube\.com\/watch\?v=([a-zA-Z0-9_-]+)',
        r'(?:https?:\/\/)?(?:www\.)?youtu\.be\/([a-zA-Z0-9_-]+)',
        r'(?:https?:\/\/)?(?:www\.)?youtube\.com\/embed\/([a-zA-Z0-9_-]+)',
        r'(?:https?:\/\/)?(?:www\.)?youtube\.com\/shorts\/([a-zA-Z0-9_-]+)'
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    parsed_url = urlparse(url)
    query_params = parse_qs(parsed_url.query)
    if 'v' in query_params:
        return query_params['v'][0]
    if parsed_url.hostname == 'youtu.be':
        return parsed_url.path.lstrip('/')
    return None

@st.cache_data(show_spinner=False)
def get_video_title_cached(video_id, youtube_link):
    """Fetch video title using YouTube Data API, fallback to ID if fails."""
    fallback_title = f"Video ID: {video_id}"
    try:
        youtube = googleapiclient.discovery.build(
            "youtube", "v3", developerKey=os.getenv("YOUTUBE_API_KEY")
        )
        response = youtube.videos().list(part="snippet", id=video_id).execute()
        title = response.get("items", [{}])[0].get("snippet", {}).get("title")
        return title if title else fallback_title
    except Exception:
        return fallback_title

def get_transcript_fallback(video_id):
    """Get transcript safely, including auto-generated captions fallback."""
    try:
        return YouTubeTranscriptApi.get_transcript(video_id)
    except (NoTranscriptFound, TranscriptsDisabled):
        try:
            return YouTubeTranscriptApi.get_transcript(video_id, languages=['en'])
        except Exception:
            return None
    except Exception:
        return None

@st.cache_data(show_spinner=False)
def get_transcript_cached(video_id):
    transcript_list = get_transcript_fallback(video_id)
    if not transcript_list:
        return None
    formatted_transcript = []
    for item in transcript_list:
        timestamp = int(item["start"])
        time_str = f"[{timestamp // 3600:02d}:{(timestamp % 3600) // 60:02d}:{timestamp % 60:02d}]"
        formatted_transcript.append(f"{time_str} {item['text']}")
    return " ".join(formatted_transcript)

def format_response(response_text):
    """Clean up Gemini response formatting."""
    cleaned = re.sub(r'\n{3,}', '\n\n', response_text)
    cleaned = re.sub(r'(?m)^\s*[-•]\s*', '- ', cleaned)
    return cleaned

# ---------------- Gemini Content Generation ----------------
if "gemini_model" not in st.session_state:
    st.session_state.gemini_model = genai.GenerativeModel(
        "gemini-1.5-flash-latest",
        generation_config={'temperature':0.7, 'top_p':0.8, 'top_k':40}
    )

def generate_content(text, prompt):
    """Generate content using Gemini."""
    model = st.session_state.gemini_model
    try:
        response = model.generate_content(f"{prompt}\n\n{text}")
        return format_response(response.text)
    except Exception as e:
        st.error(f"Error generating content: {str(e)}")
        return None

def chunk_text_smarter(text, max_length=3000):
    """Split transcript into smart chunks."""
    sentences = re.split(r'(?<=[.!?]) +', text)
    chunks = []
    current = ""
    for s in sentences:
        if len(current) + len(s) <= max_length:
            current += " " + s
        else:
            chunks.append(current.strip())
            current = s
    if current:
        chunks.append(current.strip())
    return chunks

def analyze_transcript_parallel(transcript):
    """Analyze transcript in parallel."""
    chunks = chunk_text_smarter(transcript)
    chunk_analyses = []
    progress_bar = st.progress(0)
    with ThreadPoolExecutor(max_workers=3) as executor:
        future_to_chunk = {executor.submit(generate_content, chunk, CHUNK_PROMPT): chunk for chunk in chunks}
        for i, future in enumerate(as_completed(future_to_chunk)):
            result = future.result()
            if result:
                chunk_analyses.append(result)
            progress_bar.progress((i+1)/len(chunks))
    combined_analysis = "\n\n".join(chunk_analyses)
    return generate_content(combined_analysis, FINAL_PROMPT)

def generate_qa_response(question, transcript, summary):
    """Generate answer for user question."""
    prompt = f"Question: {question}\n\nSummary:\n{summary}\n\nTranscript:\n{transcript}\n\nAnswer:"
    return generate_content(prompt, "")

# ---------------- Document Generation ----------------
def create_word_document(summary, video_title, video_id, qa_history=None):
    doc = Document()
    title = doc.add_heading(f"Video Summary: {video_title or 'Untitled'}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Video Link: https://youtube.com/watch?v={video_id}")
    doc.add_paragraph()
    for line in summary.split("\n"):
        doc.add_paragraph(line.strip())
    if qa_history:
        doc.add_heading("Questions & Answers", 1)
        for qa in qa_history:
            q_para = doc.add_paragraph()
            q_para.add_run("Q: ").bold = True
            q_para.add_run(qa['question'])
            a_para = doc.add_paragraph()
            a_para.add_run("A: ").bold = True
            a_para.add_run(qa['answer'])
            doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph("Generated using YouTube Video Summarizer")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc

def get_word_binary(video_id, summary, qa_history, video_title):
    key = f"word_{video_id}"
    if key in st.session_state:
        return st.session_state[key]
    doc = create_word_document(summary, video_title, video_id, qa_history)
    bio = io.BytesIO()
    doc.save(bio)
    st.session_state[key] = bio.getvalue()
    return st.session_state[key]

def create_markdown_download(summary, video_title, video_id, qa_history=None):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    markdown_content = f"# Video Summary: {video_title}\n\nGenerated on: {timestamp}\nVideo Link: https://youtube.com/watch?v={video_id}\n\n{summary}\n"
    if qa_history:
        markdown_content += "\n## Questions & Answers\n"
        for qa in qa_history:
            markdown_content += f"**Q: {qa['question']}**\n\nA: {qa['answer']}\n\n"
    markdown_content += "\n---\nGenerated using YouTube Video Summarizer"
    return markdown_content

# ---------------- Streamlit UI ----------------
def setup_streamlit_ui():
    st.markdown("""
    <style>
    .main-title {font-size:3rem; font-weight:800; text-align:center; color: linear-gradient(120deg,#4285F4,#0F9D58);}
    .subtitle {text-align:center;color:#5f6368;font-size:1.3rem;}
    </style>
    """, unsafe_allow_html=True)
    st.markdown('<h1 class="main-title">🎯 YouTube Video Summary</h1>', unsafe_allow_html=True)
    st.markdown('<p class="subtitle">AI-Powered Video Analysis & Summarization</p>', unsafe_allow_html=True)

def show_quick_guide():
    with st.expander("ℹ️ How to Use"):
        st.markdown("""
        1. Paste any YouTube URL
        2. Click 'Generate Detailed Notes' or 'Fast Summary'
        3. Get AI-powered summary
        4. Ask questions about the content
        5. Download in Markdown or Word format
        """)

def show_footer():
    st.markdown("---")
    st.markdown("Generated using Gemini Flash AI & Streamlit")

# ---------------- Main App ----------------
def main():
    setup_streamlit_ui()
    show_quick_guide()
    
    youtube_link = st.text_input("Enter YouTube URL:")
    video_id = get_youtube_video_id(youtube_link) if youtube_link else None
    
    if video_id:
        if st.button("Process Video") or st.session_state.current_video_id != video_id:
            st.session_state.video_processed = False
            st.session_state.current_video_id = video_id
            st.session_state.current_video_title = get_video_title_cached(video_id, youtube_link)
            st.session_state.current_transcript = get_transcript_cached(video_id)
            if st.session_state.current_transcript:
                st.success(f"Transcript fetched: {st.session_state.current_video_title}")
                st.session_state.video_processed = True
            else:
                st.warning("Transcript not available for this video. You can still generate fast summary.")

        if st.session_state.video_processed or st.session_state.current_transcript:
            col1, col2 = st.columns(2)
            with col1:
                if st.button("Generate Detailed Notes") and st.session_state.current_transcript:
                    st.session_state.current_summary = analyze_transcript_parallel(st.session_state.current_transcript)
                    st.success("Detailed summary generated!")
                    st.session_state.fast_summary_generated = True

            with col2:
                if st.button("Generate Fast Summary"):
                    transcript_text = st.session_state.current_transcript or ""
                    st.session_state.current_summary = generate_content(
                        transcript_text, FAST_SUMMARY_PROMPT
                    )
                    st.success("Fast summary generated!")
                    st.session_state.fast_summary_generated = True

        if st.session_state.current_summary:
            st.subheader("📄 Video Summary")
            st.text_area("Summary", st.session_state.current_summary, height=400)

            # QA Section
            question = st.text_input("Ask a question about this video:")
            if question:
                answer = generate_qa_response(question, st.session_state.current_transcript or "", st.session_state.current_summary)
                st.session_state.qa_history.append({'question': question, 'answer': answer})
                st.success("Answer generated!")
                st.text_area("Answer", answer, height=150)

            # Download buttons
            col1, col2 = st.columns(2)
            with col1:
                word_binary = get_word_binary(video_id, st.session_state.current_summary, st.session_state.qa_history, st.session_state.current_video_title)
                st.download_button("📥 Download Word", word_binary, f"{st.session_state.current_video_title}.docx")

            with col2:
                markdown_text = create_markdown_download(st.session_state.current_summary, st.session_state.current_video_title, video_id, st.session_state.qa_history)
                st.download_button("📥 Download Markdown", markdown_text, f"{st.session_state.current_video_title}.md")

    show_footer()

# ---------------- Run App ----------------
if __name__ == "__main__":
    main()
